"""
Document Parser — Extract text and tables from .docx files.
Detects document type based on heading/title.
"""

from pathlib import Path
from docx import Document
import fitz  # PyMuPDF


# Document type detection patterns
DOC_TYPE_PATTERNS = {
    "formulir_klaim": ["formulir pengajuan klaim", "formulir klaim"],
    "kuitansi": ["kuitansi", "rincian biaya"],
    "hasil_penunjang": ["hasil pemeriksaan penunjang", "laboratorium"],
    "hasil_rontgen": ["hasil pemeriksaan radiologi", "rontgen", "x-ray"],
    "kronologi_kecelakaan": ["kronologi", "kecelakaan", "kronologis kejadian"],
    "riwayat_preexisting": ["riwayat klaim", "pre-existing", "deklarasi kesehatan"],
}


def detect_doc_type(text: str) -> str:
    """Detect document type from first few lines of text."""
    header = text[:500].lower()
    for doc_type, patterns in DOC_TYPE_PATTERNS.items():
        if any(p in header for p in patterns):
            return doc_type
    return "unknown"


def parse_docx(file_path: str | Path) -> dict:
    """
    Parse a .docx file and extract all text content.

    Returns:
        {
            "file_name": str,
            "doc_type": str,
            "raw_text": str,
            "tables": list[list[list[str]]],  # list of tables, each table is rows of cells
            "paragraphs": list[str],
        }
    """
    file_path = Path(file_path)
    doc = Document(str(file_path))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)
        tables.append(table_data)

    raw_text = "\n".join(paragraphs)
    for table in tables:
        for row in table:
            raw_text += "\n" + " | ".join(row)

    doc_type = detect_doc_type(raw_text)

    return {
        "file_name": file_path.name,
        "doc_type": doc_type,
        "raw_text": raw_text,
        "tables": tables,
        "paragraphs": paragraphs,
    }


def parse_pdf(file_path: str | Path) -> dict:
    """
    Parse a PDF file and extract text content using PyMuPDF.

    Returns same structure as parse_docx for interoperability.
    """
    file_path = Path(file_path)
    doc = fitz.open(str(file_path))

    pages_text = []
    for page in doc:
        pages_text.append(page.get_text().strip())
    doc.close()

    raw_text = "\n".join(pages_text)
    paragraphs = [p for p in raw_text.split("\n") if p.strip()]
    doc_type = detect_doc_type(raw_text)

    return {
        "file_name": file_path.name,
        "doc_type": doc_type,
        "raw_text": raw_text,
        "tables": [],
        "paragraphs": paragraphs,
    }


def parse_file(file_path: str | Path) -> dict:
    """Route to the correct parser based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        return {
            "file_name": Path(file_path).name,
            "doc_type": "unsupported",
            "raw_text": "",
            "tables": [],
            "paragraphs": [],
            "error": f"Unsupported file format: {ext}",
        }


def parse_multiple_docs(file_paths: list[str | Path]) -> list[dict]:
    """Parse multiple document files and return list of parsed results."""
    results = []
    for fp in file_paths:
        try:
            results.append(parse_file(fp))
        except Exception as e:
            results.append({
                "file_name": Path(fp).name,
                "doc_type": "error",
                "raw_text": "",
                "tables": [],
                "paragraphs": [],
                "error": str(e),
            })
    return results


def combine_docs_text(parsed_docs: list[dict]) -> str:
    """Combine all parsed documents into a single text for LLM extraction."""
    sections = []
    for doc in parsed_docs:
        if doc.get("error"):
            continue
        sections.append(f"--- DOKUMEN: {doc['file_name']} (Tipe: {doc['doc_type']}) ---")
        sections.append(doc["raw_text"])
        sections.append("")
    return "\n".join(sections)
